"""
Converts open-plan-domain.md to open-plan-domain.docx using python-docx.
The Mermaid fenced-code block is replaced by a note that the diagram
lives in open-plan-domain.mmd (Word cannot render Mermaid natively).
Run:  python3 MW/uml/build_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path(__file__).parent / "open-plan-domain.md"
OUT = Path(__file__).parent / "open-plan-domain.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_heading_color(paragraph, r, g, b):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_table_from_md(doc, header_row, rows):
    col_count = len(header_row)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Light Shading Accent 1"
    # header
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = cell_text.strip()
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    # data rows
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < col_count:
                table.rows[ri + 1].cells[ci].text = cell_text.strip()
    doc.add_paragraph()  # spacing after table


def parse_md_table(lines):
    """Return (header_row, data_rows) from a block of MD table lines."""
    header = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    data = []
    for line in lines[2:]:  # skip separator row
        if line.strip().startswith("|"):
            data.append([c.strip() for c in line.strip().strip("|").split("|")])
    return header, data


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    md_text = SRC.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    i = 0
    in_mermaid = False
    mermaid_placeholder_added = False
    table_lines: list[str] = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # ---- Mermaid fenced block ----------------------------------------
        if line.strip() == "```mermaid":
            in_mermaid = True
            i += 1
            continue
        if in_mermaid:
            if line.strip() == "```":
                in_mermaid = False
                if not mermaid_placeholder_added:
                    p = doc.add_paragraph()
                    p.style = doc.styles["Normal"]
                    run = p.add_run(
                        "[Klassendiagram: zie open-plan-domain.mmd – "
                        "te renderen via Mermaid Live Editor of sphinxcontrib-mermaid]"
                    )
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    mermaid_placeholder_added = True
            i += 1
            continue

        # ---- Generic fenced code block ------------------------------------
        if line.strip().startswith("```"):
            # skip until closing fence
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        # ---- Horizontal rule ---------------------------------------------
        if line.strip() == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        # ---- Markdown table (buffering) ----------------------------------
        if line.strip().startswith("|"):
            table_lines.append(line)
            i += 1
            # collect whole table
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 3:
                header, rows = parse_md_table(table_lines)
                add_table_from_md(doc, header, rows)
            table_lines = []
            continue

        # ---- Headings ----------------------------------------------------
        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            p = doc.add_heading(text, level=level)
            set_heading_color(p, 0x2E, 0x74, 0xB5)  # Word-blue
            i += 1
            continue

        # ---- Blockquote --------------------------------------------------
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            # strip inline backticks
            text = re.sub(r"`([^`]+)`", r"\1", text)
            p = doc.add_paragraph(style="Quote")
            p.add_run(text)
            i += 1
            continue

        # ---- Blank line --------------------------------------------------
        if line.strip() == "":
            i += 1
            continue

        # ---- Normal paragraph (with inline formatting) -------------------
        # Strip leading list markers
        list_match = re.match(r"^[-*]\s+(.*)", line)
        is_list = bool(list_match)
        text = list_match.group(1) if list_match else line

        p = doc.add_paragraph(style="List Bullet" if is_list else "Normal")

        # Split on bold (**text**) and inline code (`text`)
        parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            elif part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part)

        i += 1

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    build()
